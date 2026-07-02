#!/usr/bin/env python3
"""生成"7轴机械臂MATLAB建模与仿真"课程设计报告Word文档"""

import os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

BASE = r'E:\Technical_support\codes\Git_Projects\GitHub\7-axis-robotic-arm'
TEMPLATE = os.path.join(BASE, 'doc', '基于funasr实现语音实时识别.docx')
OUTPUT   = os.path.join(BASE, 'doc', '7轴机械臂matlab建模与仿真-机器人231-谢毅捷.docx')
RES_PNG  = os.path.join(BASE, 'res', 'png')

# ── Load template to inherit styles ──────────────────────────
doc = Document(TEMPLATE)

# ── Remove all content from the template ─────────────────────
# Remove all paragraphs
for p in doc.paragraphs:
    p._element.getparent().remove(p._element)
# Remove all tables
for t in doc.tables:
    t._element.getparent().remove(t._element)

# ── Helper functions ─────────────────────────────────────────
def add_para(text, style='Normal', bold=False, size=None, alignment=None, space_after=None, space_before=None, first_line_indent=None, font_name=None):
    """Add a paragraph with specified formatting."""
    p = doc.add_paragraph(text, style=style)
    if bold or size or font_name:
        for run in p.runs:
            if bold:
                run.bold = bold
            if size:
                run.font.size = Pt(size)
            if font_name:
                run.font.name = font_name
                run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    if alignment is not None:
        p.alignment = alignment
    pf = p.paragraph_format
    if space_after is not None:
        pf.space_after = Pt(space_after)
    if space_before is not None:
        pf.space_before = Pt(space_before)
    if first_line_indent is not None:
        pf.first_line_indent = Cm(first_line_indent)
    return p

def add_heading(text, level=1):
    """Add a heading."""
    h = doc.add_heading(text, level=level)
    return h

def add_image(img_name, width_inches=5.5, caption='', alignment=WD_ALIGN_PARAGRAPH.CENTER):
    """Add an image with optional caption."""
    img_path = os.path.join(RES_PNG, img_name)
    if os.path.exists(img_path):
        p_img = doc.add_paragraph()
        p_img.alignment = alignment
        run = p_img.add_run()
        run.add_picture(img_path, width=Inches(width_inches))
        p_img.paragraph_format.space_before = Pt(6)
        p_img.paragraph_format.space_after = Pt(3)
    if caption:
        p_cap = doc.add_paragraph(caption)
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p_cap.runs:
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(100, 100, 100)

def add_bullet(text, level=0):
    """Add a bullet point."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.0 + level * 0.7)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(1)
    run = p.add_run('• ' + text)
    set_run_font(run, '宋体', 11)
    return p

def add_page_break():
    """Add a page break."""
    doc.add_page_break()

def set_run_font(run, name='宋体', size=12, bold=False, color=None):
    """Set font for a run."""
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

# ═══════════════════════════════════════════════════════════════
#  COVER PAGE (封面)
# ═══════════════════════════════════════════════════════════════

# Empty lines for spacing
for _ in range(6):
    add_para('', space_after=12)

add_para('7轴机械臂MATLAB建模与仿真', style='Normal', bold=True, size=22,
         alignment=WD_ALIGN_PARAGRAPH.CENTER, font_name='黑体')
add_para('', space_after=6)

# Subtitle line
p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_sub = p_sub.add_run('——基于Robotics System Toolbox的运动学仿真与轨迹规划')
set_run_font(run_sub, '宋体', 16)

for _ in range(6):
    add_para('', space_after=12)

# Info lines
info_items = [
    ('学       院：', '自动化与电气工程学院'),
    ('专       业：', '机器人工程'),
    ('姓       名：', '谢毅捷'),
    ('班       级：', '机器人231班'),
    ('学       号：', '2023303050130'),
    ('日       期：', '2026.07.02'),
]
for label, value in info_items:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_l = p.add_run(label)
    set_run_font(run_l, '宋体', 16, bold=True)
    run_v = p.add_run(value)
    set_run_font(run_v, '宋体', 16)

add_page_break()

# ═══════════════════════════════════════════════════════════════
#  ABSTRACT (摘要)
# ═══════════════════════════════════════════════════════════════
add_para('摘  要', style='Normal', bold=True, size=16,
         alignment=WD_ALIGN_PARAGRAPH.CENTER, font_name='黑体', space_after=12)

abstract_text = (
    '本文基于MATLAB Robotics System Toolbox，设计并实现了一套完整的七自由度机械臂运动学仿真与轨迹规划系统。'
    '机械臂采用Z-Y-Z-Y-Z-Y-Z七旋转关节串联构型，工作半径约为0.87米，包含底座旋转（腰）、肩部俯仰、大臂横滚（冗余自由度）、'
    '肘部俯仰、小臂横滚、腕部俯仰和末端横滚共七个关节。系统围绕刚体树（RigidBodyTree）模型，'
    '建立了包含底座、七段连杆和末端执行器的精确运动学模型，每段关节均设置了合理的运动范围约束。'
    '在正运动学方面，基于齐次变换矩阵实现了关节空间到笛卡尔空间的映射；在逆运动学方面，'
    '采用广义逆运动学（Generalized Inverse Kinematics, GIK）求解器，通过数值迭代方法实现末端位姿到关节角度的求解，'
    '支持仅位置约束和位姿全约束两种模式。此外，实现了基于几何雅可比矩阵的速度运动学分析，验证了机械臂在工作空间内的灵活性。'
)
add_para(abstract_text, first_line_indent=0.74, space_after=6)

abstract_text2 = (
    '在轨迹规划方面，系统集成了四种主流轨迹生成方法：梯形速度剖面法（trapveltraj）、三次多项式插值（cubicpolytraj）、'
    '五次多项式插值（quinticpolytraj）和B样条曲线插值（bsplinepolytraj）。用户可通过图形用户界面（GUI）自由编辑路径点，'
    '配置IK求解参数和轨迹参数，系统自动完成路径点逆解、关节空间轨迹生成和三维可视化动画播放，'
    '并弹出各关节的角度、角速度和角加速度曲线供分析。交互界面基于MATLAB App Designer框架开发，'
    '采用双标签页设计——"探索"页面提供FK/IK滑块交互式运动学探索，"规划"页面提供完整的轨迹规划工作流。'
    '测试结果表明，系统运行稳定，正运动学计算精确，逆运动学求解成功率高，轨迹规划平滑连续，具有良好的教学演示和工程仿真价值。'
)
add_para(abstract_text2, first_line_indent=0.74, space_after=12)

# Keywords
p_kw = doc.add_paragraph()
p_kw.paragraph_format.first_line_indent = Cm(0.74)
run_kw_label = p_kw.add_run('关键词：')
set_run_font(run_kw_label, '宋体', 12, bold=True)
run_kw = p_kw.add_run('七自由度机械臂；MATLAB；运动学仿真；轨迹规划；Robotics System Toolbox；GUI')
set_run_font(run_kw, '宋体', 12)

add_page_break()

# ═══════════════════════════════════════════════════════════════
#  TABLE OF CONTENTS (目录 placeholder)
# ═══════════════════════════════════════════════════════════════
add_para('目  录', style='Normal', bold=True, size=16,
         alignment=WD_ALIGN_PARAGRAPH.CENTER, font_name='黑体', space_after=12)
add_para('（请在Word中右键此处 → 更新域 → 更新整个目录，以自动生成目录）',
         alignment=WD_ALIGN_PARAGRAPH.CENTER, size=10)
add_para('', space_after=6)

# Manual TOC entries
toc_entries = [
    ('摘  要', 'I'),
    ('目  录', 'II'),
    ('1  项目背景与目标', '1'),
    ('    1.1  项目背景', '1'),
    ('    1.2  项目目标', '2'),
    ('2  机械臂运动学建模', '3'),
    ('    2.1  机械臂结构设计', '3'),
    ('    2.2  DH参数与坐标变换', '4'),
    ('    2.3  正运动学分析', '5'),
    ('    2.4  逆运动学分析', '6'),
    ('    2.5  雅可比矩阵与速度运动学', '7'),
    ('3  系统设计与实现', '8'),
    ('    3.1  系统总体架构', '8'),
    ('    3.2  交互式GUI界面设计', '9'),
    ('    3.3  轨迹规划算法设计', '11'),
    ('4  实验与测试', '13'),
    ('    4.1  测试环境与方案', '13'),
    ('    4.2  模型结构验证', '13'),
    ('    4.3  正运动学测试', '14'),
    ('    4.4  逆运动学测试', '15'),
    ('    4.5  轨迹规划测试', '16'),
    ('5  总结与展望', '18'),
    ('    5.1  工作总结', '18'),
    ('    5.2  不足与展望', '18'),
    ('参考文献', '19'),
    ('致  谢', '20'),
]
for entry, page in toc_entries:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(0)
    run_e = p.add_run(entry)
    set_run_font(run_e, '宋体', 12)
    # Add dots and page number
    dots = '·' * max(2, 50 - len(entry))
    run_d = p.add_run(f' {dots} ')
    set_run_font(run_d, '宋体', 10)
    run_p = p.add_run(page)
    set_run_font(run_p, '宋体', 12)

add_page_break()

# ╔══════════════════════════════════════════════════════════════╗
# ║  CHAPTER 1 — 项目背景与目标                                  ║
# ╚══════════════════════════════════════════════════════════════╝

add_heading('1  项目背景与目标', level=1)

add_heading('1.1  项目背景', level=2)

ch1_texts = [
    ('工业背景',
     '随着"中国制造2025"和工业4.0战略的深入推进，工业机器人作为智能制造的核心装备，'
     '在汽车制造、电子装配、仓储物流、医疗手术等领域的应用日益广泛。根据国际机器人联合会（IFR）统计，'
     '2024年全球工业机器人装机量已突破60万台，其中中国连续多年位居全球最大市场。'
     '多轴串联机械臂因其工作空间大、灵活性高、可重构性强等优势，成为工业机器人的主流构型。'),
    ('技术背景',
     '机械臂运动学是机器人控制的基础理论，涉及正运动学（Forward Kinematics, FK）、'
     '逆运动学（Inverse Kinematics, IK）、速度运动学和动力学四大核心问题。其中，逆运动学求解是机械臂轨迹规划与控制的关键环节，'
     '对于七自由度冗余机械臂，其逆解存在无穷多组解，需要借助数值优化方法进行求解。'
     'MATLAB Robotics System Toolbox提供了丰富的机器人建模、运动学分析和轨迹规划函数，'
     '为快速搭建机械臂仿真系统提供了强大的工具支持。'),
    ('七自由度机械臂的优势',
     '相比传统的六自由度机械臂，七自由度机械臂在末端执行器位姿固定的情况下，'
     '仍保留一个冗余自由度，使得机械臂可以在不改变末端位姿的前提下调整肘部等中间关节的构型。'
     '这一特性赋予了七轴机械臂出色的避障能力——当机械臂在狭小或复杂环境中作业时，'
     '可以通过改变冗余关节的角度来绕开障碍物，而无需改变末端工具的位置和姿态。'
     '此外，冗余自由度还可以用于优化关节力矩分布、避免关节极限和奇异位形，'
     '提高机械臂的运动平滑性和操作灵活性，因此七轴机械臂在空间站维护、微创手术、'
     '人机协作等高端应用场景中具有不可替代的优势。'),
]
for title, text in ch1_texts:
    add_para(title, bold=True, size=12, space_before=6, space_after=3)
    add_para(text, first_line_indent=0.74, space_after=6)

add_heading('1.2  项目目标', level=2)

add_para('本项目的总体目标是基于MATLAB Robotics System Toolbox，构建一套功能完整、'
         '交互友好的七自由度机械臂运动学仿真与轨迹规划系统，具体目标如下：', first_line_indent=0.74, space_after=6)

goals = [
    ('机械臂建模目标：', '建立包含底座、七段旋转连杆和末端执行器的精确刚体树模型，'
     '为每段连杆添加符合实际尺寸的可视化几何体（圆柱体关节外壳+长方体连杆），'
     '设置合理的关节限位角度，使模型在全工作空间内具有逼真的三维显示效果。'),
    ('运动学分析目标：', '实现正运动学的齐次变换矩阵计算，将关节角度映射为末端执行器的三维位置和RPY姿态角；'
     '实现基于广义逆运动学求解器的逆运动学算法，支持仅位置约束和位姿全约束两种模式；'
     '实现几何雅可比矩阵的计算与秩分析，验证机械臂在工作空间内的灵活性和奇异性特征。'),
    ('轨迹规划目标：', '集成梯形速度剖面法、三次多项式、五次多项式和B样条共四种关节空间轨迹规划方法，'
     '支持用户自定义路径点序列、总运动时间、每段采样点数和峰值速度等参数，'
     '自动生成各关节的角度、角速度和角加速度连续曲线，并在三维视图中实时播放机械臂运动动画。'),
    ('GUI开发目标：', '基于MATLAB App Designer框架开发统一的图形用户界面，'
     '将运动学探索（FK/IK滑块交互）和轨迹规划两大功能集成在同一界面的两个标签页中，'
     '提供直观的滑块拖动、实时三维视图更新、末端位姿显示和工作状态反馈。'),
    ('验证测试目标：', '编写系统化的测试脚本，覆盖模型结构完整性、正运动学精度、'
     '逆运动学求解成功率、雅可比矩阵秩分析、自碰撞检测和轨迹生成一致性共八个测试项，'
     '确保仿真系统的正确性和可靠性。'),
]
for title, text in goals:
    add_para(title + text, first_line_indent=0.74, space_after=4)

add_page_break()

# ╔══════════════════════════════════════════════════════════════╗
# ║  CHAPTER 2 — 机械臂运动学建模                               ║
# ╚══════════════════════════════════════════════════════════════╝

add_heading('2  机械臂运动学建模', level=1)

add_heading('2.1  机械臂结构设计', level=2)

add_para('本项目的七自由度机械臂采用Z-Y-Z-Y-Z-Y-Z的全旋转关节串联构型，即相邻关节的旋转轴交替沿Z轴和Y轴方向。'
         '在零位构型下（所有关节角均为0°），机械臂竖直向上沿世界坐标系+Z轴方向伸展，总伸长度为0.87米。'
         '机械臂由底座（pedestal）、七段旋转连杆（link1-link7）和末端执行器（ee）共9个刚体组成，'
         '各连杆之间通过旋转关节（revolute joint）依次串联连接。', first_line_indent=0.74, space_after=6)

add_para('各关节的布局与功能如下：', bold=True, space_after=4)

joint_desc = [
    'J1 — 底座旋转（腰关节），绕Z轴旋转，运动范围 ±170°，负责机械臂的水平回转',
    'J2 — 肩部俯仰，绕Y轴旋转，运动范围 ±120°，控制大臂的前后摆动',
    'J3 — 大臂横滚（冗余自由度），绕Z轴旋转，运动范围 ±170°，在大臂方向上提供自转自由度',
    'J4 — 肘部俯仰，绕Y轴旋转，运动范围 ±120°，控制小臂的屈伸',
    'J5 — 小臂横滚，绕Z轴旋转，运动范围 ±170°，在小臂方向上提供自转自由度',
    'J6 — 腕部俯仰，绕Y轴旋转，运动范围 ±120°，控制腕部的俯仰姿态',
    'J7 — 末端横滚，绕Z轴旋转，运动范围 ±175°，控制末端工具的自转',
]
for jd in joint_desc:
    add_bullet(jd)

add_para('', space_after=4)

add_para('在各连杆之间，通过R_x(-90°)和R_x(+90°)的固定坐标变换交替转换坐标系，'
         '使得相邻两段变换互相抵消。这种设计巧妙地保证了在零位时各奇数段关节的本体坐标系与世界坐标系保持一致，'
         '简化了运动学模型的推导和验证。各段连杆的几何尺寸参数如表2-1所示。', first_line_indent=0.74, space_after=8)

# Table 2-1: Link dimensions
add_para('表2-1  机械臂各段尺寸参数', bold=True, size=10,
         alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
table = doc.add_table(rows=13, cols=3, style='Table Grid')
table.alignment = WD_ALIGN_PARAGRAPH.CENTER
headers = ['参数名称', '符号', '数值（m）']
data = [
    ['底座高度', 'baseH', '0.12'],
    ['底座半径', 'baseR', '0.10'],
    ['J1→J2 偏距', 'L0', '0.05'],
    ['大臂长度（J2→J3）', 'L1', '0.28'],
    ['J3→J4 偏距', 'L2', '0.04'],
    ['小臂长度（J4→J5）', 'L3', '0.22'],
    ['J5→J6 偏距', 'L4', '0.04'],
    ['腕部长度（J6→J7）', 'L5', '0.06'],
    ['J7→末端法兰', 'L6', '0.06'],
    ['关节外壳半径', 'jR', '0.040'],
    ['大臂方盒宽度', 'bW', '0.060'],
    ['中段方盒宽度', 'bW2', '0.050'],
]
for j, h in enumerate(headers):
    cell = table.rows[0].cells[j]
    cell.text = h
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.bold = True
            r.font.size = Pt(9)
for i, row_data in enumerate(data):
    for j, val in enumerate(row_data):
        cell = table.rows[i+1].cells[j]
        cell.text = val
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.size = Pt(9)

add_para('', space_after=6)

add_heading('2.2  DH参数与坐标变换', level=2)

add_para('本模型采用RigidBodyTree框架下的固定变换（setFixedTransform）机制来描述相邻关节之间的空间关系，'
         '而非传统DH参数法。每个关节的位姿由两部分构成：', first_line_indent=0.74, space_after=6)

add_para('（1）固定变换（Fixed Transform）：描述父关节坐标系到当前关节旋转轴坐标系的刚性变换，'
         '包括平移和旋转两部分。平移量由连杆长度和偏距决定，旋转量采用R_x(-90°)和R_x(+90°)交替使用，'
         '实现Z轴和Y轴旋转方向的切换。以J2关节为例，其固定变换为：沿Z轴平移L0（0.05 m），再绕X轴旋转-90°，'
         '这使得J2在零位时的本体Y轴与世界-Z轴对齐，本体Z轴与世界Y轴对齐，从而实现了绕世界Y轴的俯仰运动。'
         '对于J3关节，固定变换为沿Y轴平移-L1（-0.28 m），再绕X轴旋转+90°，'
         '这一变换与前一段的R_x(-90°)互相抵消，使得J3零位时的本体坐标系重新与世界坐标系对齐。', first_line_indent=0.74, space_after=4)

add_para('（2）关节变量（Joint Variable）：旋转关节的角度变量q_i，叠加在固定变换之上，'
         '绕当前本体坐标系的旋转轴（Z轴或Y轴）转动。Z轴关节（J1、J3、J5、J7）实现横滚运动，'
         'Y轴关节（J2、J4、J6）实现俯仰运动，两者交替排列形成了Z-Y-Z-Y-Z-Y-Z的关节构型序列。', first_line_indent=0.74, space_after=4)

add_para('这种固定变换+关节变量的两层结构可表示为：', first_line_indent=0.74, space_after=4)

# Formula: T_i = T_fixed * R_axis(q_i)
p_formula = doc.add_paragraph()
p_formula.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_f = p_formula.add_run('T_i(q_i) = T_fixed,i · R_axis(q_i)')
set_run_font(run_f, 'Times New Roman', 12, bold=False)

add_para('其中T_fixed,i为第i关节的固定齐次变换矩阵，R_axis(q_i)为绕当前关节轴旋转q_i弧度的旋转矩阵。'
         '末端执行器相对于世界坐标系的总变换为所有关节变换的连乘积：', first_line_indent=0.74, space_after=4)

p_formula2 = doc.add_paragraph()
p_formula2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_f2 = p_formula2.add_run('T_ee(q₁,...,q₇) = T₁(q₁)·T₂(q₂)·...·T₇(q₇)·T_fixed,ee')
set_run_font(run_f2, 'Times New Roman', 12, bold=False)

add_para('', space_after=4)

add_heading('2.3  正运动学分析', level=2)

add_para('正运动学（Forward Kinematics, FK）解决的是"已知各关节角度，求末端执行器位姿"的问题。'
         '在MATLAB Robotics System Toolbox中，getTransform函数通过遍历刚体树结构，'
         '从根节点（base）到目标体（ee）依次计算各关节的齐次变换矩阵并连乘，得到末端执行器坐标系相对于世界坐标系的4×4齐次变换矩阵T：',
         first_line_indent=0.74, space_after=6)

# T matrix description
p_mat = doc.add_paragraph()
p_mat.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_mat = p_mat.add_run('     ┌                  ┐\n'
                        '     │ R₃×₃     P₃×₁   │\n'
                        'T =  │                    │\n'
                        '     │ 0  0  0    1      │\n'
                        '     └                  ┘')
set_run_font(run_mat, 'Times New Roman', 11)

add_para('其中R₃×₃为3×3旋转矩阵，描述末端姿态；P₃×₁ = [x, y, z]ᵀ 为3×1平移向量，描述末端位置。'
         '从旋转矩阵中可以进一步提取RPY欧拉角（ZYX顺序）：偏航角Yaw（绕Z轴）、俯仰角Pitch（绕Y轴）、'
         '横滚角Roll（绕X轴），即R = Rz(Yaw)·Ry(Pitch)·Rx(Roll)。',
         first_line_indent=0.74, space_after=6)

add_para('在零位构型（q=[0,0,0,0,0,0,0]）下，正运动学计算结果为末端位置（0, 0, 0.87）米，'
         '即机械臂竖直向上完全伸展时，末端位于世界坐标系原点正上方0.87米处，与理论设计值吻合。'
         '在GUI的"探索"标签页中，用户拖动六个关节滑块时，系统实时调用getTransform计算末端位姿，'
         '并在三维视图中同步更新机械臂构型和末端位置显示，实现了直观的FK交互探索。',
         first_line_indent=0.74, space_after=8)

add_image('探索界面.png', width_inches=5.0, caption='图2-1  FK/IK探索界面（GUI"探索"标签页）')

add_heading('2.4  逆运动学分析', level=2)

add_para('逆运动学（Inverse Kinematics, IK）解决的是"已知末端执行器目标位姿，反求各关节角度"的问题。'
         '对于七自由度冗余机械臂，给定末端位姿存在无穷多组可行关节角解（冗余度为1），'
         '因此需要借助数值迭代方法进行求解。', first_line_indent=0.74, space_after=6)

add_para('本系统采用MATLAB Robotics System Toolbox提供的广义逆运动学求解器（generalizedInverseKinematics, GIK）。'
         'GIK基于非线性优化方法，通过定义约束条件和代价函数，在关节空间中搜索满足末端位姿约束的最优关节角度。'
         '求解器支持以下关键配置：', first_line_indent=0.74, space_after=4)

ik_config = [
    '约束类型：支持两种约束模式——"仅位置约束"（constraintPositionTarget）只需末端位置到达目标点，'
    '"位置+姿态约束"（constraintPoseTarget）需同时满足末端位置和姿态的要求',
    '关节限位约束（constraintJointBounds）：确保求解结果在各关节的允许运动范围内',
    '最大迭代次数：默认400次，GUI中可在10-2000范围内调整',
    '最大求解时间：1.0秒，防止过长时间的迭代搜索',
    '初始猜测：使用当前机械臂构型作为迭代起点，利用邻近构型连续性加速收敛',
]
for item in ik_config:
    add_bullet(item)

add_para('', space_after=4)
add_para('求解过程的数学形式为：以当前关节角q_curr为初始值，在关节空间中进行迭代搜索，'
         '寻找满足末端位姿约束T_target且关节角度增量最小的解q_sol。'
         '在GUI逆运动学模式下，用户拖动XYZ/RPY滑块设定目标位姿，系统自动调用GIK求解器，'
         '若求解成功（Status="success"），则通过三次多项式插值生成从当前构型到目标构型的平滑过渡动画（25帧），'
         '同时更新关节滑块和末端位姿显示；若求解失败，界面会显示具体失败状态（如"unreachable"或"iteration limit exceeded"）'
         '并以橙色文字提示用户。', first_line_indent=0.74, space_after=8)

add_image('逆运动学演示.gif', width_inches=5.0, caption='图2-2  逆运动学求解演示')

add_heading('2.5  雅可比矩阵与速度运动学', level=2)

add_para('几何雅可比矩阵（Geometric Jacobian）J(q) ∈ ℝ⁶ˣ⁷ 建立了关节空间速度与操作空间速度之间的线性映射关系：',
         first_line_indent=0.74, space_after=4)

p_j = doc.add_paragraph()
p_j.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_j = p_j.add_run('[v; ω] = J(q) · q̇')
set_run_font(run_j, 'Times New Roman', 12)

add_para('其中v为末端线速度（3×1），ω为末端角速度（3×1），q̇为关节角速度（7×1）。'
         '对于七自由度机械臂，J(q)是一个6×7矩阵，其秩最大为6。当rank(J)=6时，机械臂处于非奇异构型，'
         '末端可以在六个自由度方向上自由运动；当rank(J)<6时，机械臂处于奇异位形，某些方向上的运动能力丧失。',
         first_line_indent=0.74, space_after=6)

add_para('在测试脚本中，选取非零位构型q_bent = [0, -30°, -60°, 0, -30°, 0, 0]，'
         '通过geometricJacobian函数计算雅可比矩阵并验证其秩为6（满行秩），'
         '确认该构型下机械臂具有完整的六自由度运动能力。雅可比矩阵的分析为进一步的'
         '速度控制、力控制和奇异性规避提供了理论基础。', first_line_indent=0.74, space_after=6)

add_page_break()

# ╔══════════════════════════════════════════════════════════════╗
# ║  CHAPTER 3 — 系统设计与实现                                  ║
# ╚══════════════════════════════════════════════════════════════╝

add_heading('3  系统设计与实现', level=1)

add_heading('3.1  系统总体架构', level=2)

add_para('本系统的软件架构遵循模块化设计原则，以MATLAB脚本和函数为基本单元，'
         '通过统一的图形用户界面（GUI）进行集成。系统由以下核心模块组成：', first_line_indent=0.74, space_after=6)

modules = [
    ('刚体树模型层（build_7axis_arm.m）：', '定义机械臂的物理结构和运动学参数，'
     '包括各刚体的几何尺寸、关节类型、旋转轴方向、运动范围限制和可视化几何体。'
     '该模块采用RigidBodyTree数据结构，是整个系统的基础数据层。'),
    ('运动学计算层：', '依托MATLAB Robotics System Toolbox的内置函数实现运动学核心算法。'
     '正运动学使用getTransform函数，逆运动学使用generalizedInverseKinematics求解器，'
     '雅可比矩阵使用geometricJacobian函数，碰撞检测使用checkCollision函数。'
     '该层为上层GUI和轨迹规划提供计算服务。'),
    ('轨迹规划层：', '集成MATLAB Robotics System Toolbox提供的四种关节空间轨迹生成函数——'
     'trapveltraj（梯形速度）、cubicpolytraj（三次多项式）、quinticpolytraj（五次多项式）、'
     'bsplinepolytraj（B样条），根据用户设置的路径点序列和运动参数生成连续平滑的关节轨迹。'),
    ('GUI交互层（robot_arm_gui.m）：', '基于MATLAB uifigure和App Designer组件框架开发的图形界面，'
     '采用双标签页布局——"探索"标签页集成FK/IK滑块交互功能，"规划"标签页集成路径点编辑、'
     'IK求解、轨迹生成和动画播放功能。该层负责用户输入捕获、参数传递、计算触发和结果可视化。'),
    ('测试验证层（run_tests.m）：', '系统化的自动化测试脚本，覆盖模型结构、正运动学精度、'
     '逆运动学求解、雅可比矩阵秩、关节限位、自碰撞检测和轨迹生成一致性共8个测试项，'
     '确保系统各模块的正确性和可靠性。'),
]
for title, text in modules:
    add_para(title + text, first_line_indent=0.74, space_after=5)

add_para('系统的工作流程如下：用户启动robot_arm_gui后，系统首先调用build_7axis_arm构建刚体树模型，'
         '然后在三维坐标系中渲染零位构型的机械臂。用户可在"探索"页面通过FK/IK滑块交互式探索机械臂运动学特性，'
         '或切换到"规划"页面编辑路径点、选择轨迹方法、运行动画仿真。所有计算和渲染均在MATLAB主线程中完成，'
         '动画播放通过递增的代数标记（animGen）实现安全的动画中断和并发控制。', first_line_indent=0.74, space_after=8)

add_heading('3.2  交互式GUI界面设计', level=2)

add_para('系统GUI基于MATLAB uifigure框架开发，主窗口尺寸为1300×780像素，'
         '采用左右分栏布局：左侧为共享的三维视图区域（占3/5宽度），右侧为功能标签页区域（占2/5宽度）。'
         '界面包含两个标签页：', first_line_indent=0.74, space_after=6)

add_para('3.2.1  "探索"标签页（Explore / 探索）', bold=True, size=12, space_after=4)

add_para('该标签页提供交互式运动学探索功能，包含FK和IK两种子模式，通过一对互斥按钮切换：', first_line_indent=0.74, space_after=4)

add_para('FK模式：提供6个关节角度滑块（J1-J6，J7固定为0），每个滑块的范围由对应关节的PositionLimits决定。'
         '用户拖动滑块时，系统使用防抖定时器（80ms延迟）减少实时更新频率，避免过度渲染。'
         '每次更新：读取6个滑块值作为目标关节角 → 使用三次多项式插值生成20-30帧平滑过渡动画 → '
         '更新三维视图中的机械臂构型 → 计算并显示末端执行器的XYZ位置和RPY姿态角。'
         '动画生成代数（animGen）标记机制确保快速连续拖动时旧动画立即中止，只播放最新动画。', first_line_indent=0.74, space_after=4)

add_para('IK模式：提供6个目标位姿滑块（XYZ位置 + RPY姿态角），用户拖动滑块设定末端目标位姿，'
         '系统自动调用GIK求解器计算对应的关节角度。求解成功后，通过25帧的平滑过渡动画从当前构型移动到目标构型，'
         '同时反向同步FK滑块和末端位姿显示。若目标位姿不可达，界面以橙色文字提示具体原因。'
         'IK模式与FK模式共享同一组状态管理变量，切换时自动同步滑块值和当前构型。', first_line_indent=0.74, space_after=6)

add_para('3.2.2  "规划"标签页（Plan / 轨迹规划）', bold=True, size=12, space_after=4)

add_para('该标签页提供完整的轨迹规划工作流，界面上到下分为六个功能区块：', first_line_indent=0.74, space_after=4)

plan_blocks = [
    '路径点编辑区（Waypoints）：可编辑的数据表格（uitable），每行表示一个三维路径点(X,Y,Z)，'
    '支持动态添加行、删除末行和重置为演示路径点（8个预设点的pick-and-place轨迹）',
    'IK设置区（IK Settings）：下拉菜单选择约束模式（位置/位姿），数值输入框设置最大迭代次数（10-2000）',
    '轨迹方法区（Trajectory Method）：下拉菜单从4种方法中选择，设置总运动时间和每段采样点数',
    '高级参数区（Advanced Parameters）：峰值速度设置（仅梯形速度法有效，0表示自动）',
    '运行控制区：绿色"Run Simulation"按钮触发完整流程，红色"Stop"按钮中断动画播放',
    '状态栏：显示当前操作进度（求解IK中/生成轨迹中/播放中）和最终统计信息（路径点数、帧数、时长、末端路径长度）',
]
for block in plan_blocks:
    add_bullet(block)

add_para('', space_after=4)
add_para('点击Run后，系统依次执行：①对每个路径点调用GIK求逆解 → ②若某个路径点IK失败则中止并报错 → '
         '③根据选定的轨迹方法生成关节空间轨迹 → ④预计算末端执行器三维路径 → ⑤在三维视图中绘制绿色路径线和红色路径点标记 → '
         '⑥逐帧播放机械臂运动动画，标题栏同步显示当前时间 → ⑦播放完成后弹出关节角度/角速度/角加速度曲线图窗口 → '
         '⑧状态栏显示轨迹统计摘要（路径点数、总帧数、总时长、末端行程距离）', first_line_indent=0.74, space_after=8)

add_image('轨迹规划界面.png', width_inches=5.0, caption='图3-1  轨迹规划界面（GUI"规划"标签页）')

add_heading('3.3  轨迹规划算法设计', level=2)

add_para('轨迹规划的目标是在关节空间中生成一条连接多个路径点的平滑轨迹，使得机械臂各关节能够协调运动，'
         '末端执行器依次经过预设的笛卡尔空间路径点。本系统集成了以下四种关节空间轨迹规划方法：',
         first_line_indent=0.74, space_after=6)

add_para('3.3.1  梯形速度剖面法（trapveltraj）', bold=True, size=12, space_after=4)
add_para('梯形速度剖面法是最经典的轨迹规划方法之一，其核心思想是将运动过程分为加速段、匀速段和减速段三个阶段，'
         '速度曲线呈梯形形状。该方法具有以下特点：', first_line_indent=0.74, space_after=4)
trapvel_features = [
    '速度曲线连续但加速度存在阶跃（在加速/匀速和匀速/减速的切换点处），会产生一定的冲击（jerk）',
    '参数调节简单直观，仅需设置总运动时间和峰值速度',
    '计算效率高，适合实时控制和简单轨迹',
    'MATLAB函数签名：[q, qd, qdd, t] = trapveltraj(waypoints, numSamples, EndTime, T, PeakVelocity, v)',
]
for f in trapvel_features:
    add_bullet(f)

add_para('', space_after=4)

add_para('3.3.2  三次多项式插值（cubicpolytraj）', bold=True, size=12, space_after=4)
add_para('三次多项式在每段路径点之间构造形如 q(t)=a₀+a₁t+a₂t²+a₃t³ 的多项式轨迹。'
         '其系数由路径点位置约束和路径点处速度连续性条件确定。'
         '三次多项式保证位置和速度的C²连续性（二阶可导），加速度曲线为分段线性。'
         '与梯形速度法相比，其加速度无阶跃变化，运动更为平滑。', first_line_indent=0.74, space_after=6)

add_para('3.3.3  五次多项式插值（quinticpolytraj）', bold=True, size=12, space_after=4)
add_para('五次多项式在每段路径点之间构造形如 q(t)=a₀+a₁t+a₂t²+a₃t³+a₄t⁴+a₅t⁵ 的多项式轨迹。'
         '其系数由路径点位置、速度和加速度的连续性条件联合确定。'
         '五次多项式保证位置、速度和加速度的C⁴连续性（四阶可导），加速度曲线为三次多项式曲线，光滑无尖点，'
         '产生的冲击（jerk）最小，适合对运动平滑性要求较高的精密操作场景。'
         '但其计算量较三次多项式大，且参数调节的直观性不如梯形速度法。', first_line_indent=0.74, space_after=6)

add_para('3.3.4  B样条曲线插值（bsplinepolytraj）', bold=True, size=12, space_after=4)
add_para('B样条（B-Spline）是一种基于控制点的参数化曲线表示方法，通过调整控制点和节点向量可以灵活地控制轨迹形状。'
         'B样条具有局部支撑性——修改一个控制点仅影响其邻近区域的曲线形状，而不会波及整条轨迹，'
         '这使其在交互式轨迹编辑和局部优化场景中具有独特优势。B样条轨迹的平滑性取决于阶次（degree），'
         '高阶B样条可保证足够高的连续性阶数。', first_line_indent=0.74, space_after=8)

# Comparison table
add_para('表3-1  四种轨迹规划方法对比', bold=True, size=10,
         alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
table2 = doc.add_table(rows=5, cols=5, style='Table Grid')
table2.alignment = WD_ALIGN_PARAGRAPH.CENTER
headers2 = ['方法', '连续性', '加速度平滑度', '计算复杂度', '适用场景']
data2 = [
    ['梯形速度', '速度C⁰连续', '存在阶跃', '低', '快速点到点运动'],
    ['三次多项式', '加速度C²连续', '分段线性', '中', '一般轨迹规划'],
    ['五次多项式', '加加速度C⁴连续', '光滑曲线', '较高', '精密操作/高平滑'],
    ['B样条', '可调（与阶次相关）', '高阶光滑', '较高', '交互编辑/局部优化'],
]
for j, h in enumerate(headers2):
    cell = table2.rows[0].cells[j]
    cell.text = h
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.bold = True
            r.font.size = Pt(9)
for i, row_data in enumerate(data2):
    for j, val in enumerate(row_data):
        cell = table2.rows[i+1].cells[j]
        cell.text = val
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.size = Pt(9)

add_para('', space_after=6)
add_para('轨迹规划完成后，系统弹出一个新窗口展示6个关节的角度、角速度和角加速度随时间变化的曲线，'
         '便于进行运动学特性分析。各关节曲线使用不同颜色区分（lines(6)配色），'
         '三个子图纵向排列，横轴统一为时间（秒），纵轴分别为角度（度）、角速度（度/秒）和角加速度（度/秒²）。',
         first_line_indent=0.74, space_after=8)

# Show both joint curve images side by side conceptually
add_image('梯形速度关节曲线.png', width_inches=5.0, caption='图3-2  梯形速度剖面法关节运动曲线')
add_image('五次多项式关节曲线.png', width_inches=5.0, caption='图3-3  五次多项式法关节运动曲线')

add_page_break()

# ╔══════════════════════════════════════════════════════════════╗
# ║  CHAPTER 4 — 实验与测试                                      ║
# ╚══════════════════════════════════════════════════════════════╝

add_heading('4  实验与测试', level=1)

add_heading('4.1  测试环境与方案', level=2)

add_para('系统测试在以下环境中进行：', first_line_indent=0.74, space_after=4)

test_env = [
    '操作系统：Windows 11 Home China（64位）',
    'MATLAB版本：R2020b+',
    '关键工具箱：Robotics System Toolbox',
    '硬件配置：标准笔记本电脑',
    '测试脚本：run_tests.m',
]
for item in test_env:
    add_bullet(item)

add_para('', space_after=4)
add_para('测试方案采用自动化单元测试与交互式功能验证相结合的方式。自动化测试覆盖模型结构完整性、'
         '正运动学精度、逆运动学求解成功率、雅可比矩阵秩、关节限位合理性和轨迹生成一致性共8个测试项。'
         '每条测试用例包含明确的断言（assert），通过/失败状态以彩色文字（PASS/FAIL）输出到MATLAB命令窗口，'
         '最后汇总显示通过和失败的测试数。', first_line_indent=0.74, space_after=8)

add_heading('4.2  模型结构验证', level=2)

add_para('测试1验证刚体树模型的结构完整性：检查机械臂是否包含正确数量的刚体（9个，含底座和末端执行器），'
         '以及各刚体的名称是否按预期顺序排列——pedestal（底座）→ link1~link7（七段连杆）→ ee（末端执行器）。'
         '测试结果确认模型结构符合设计规格。', first_line_indent=0.74, space_after=4)

add_para('测试6验证各关节的限位范围设置是否合理：逐一检查6个旋转关节（J1-J6）的PositionLimits属性，'
         '确保下限小于上限，且运动范围（上限-下限）大于0.5弧度（约28.6°），为机械臂提供足够的运动空间。'
         '所有关节均通过了限位合理性检查。', first_line_indent=0.74, space_after=6)

# Table 4-1: Joint limits
add_para('表4-1  各关节运动范围', bold=True, size=10,
         alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
table3 = doc.add_table(rows=8, cols=4, style='Table Grid')
table3.alignment = WD_ALIGN_PARAGRAPH.CENTER
headers3 = ['关节', '名称', '下限（°）', '上限（°）']
data3 = [
    ['J1', '腰', '-170', '+170'],
    ['J2', '肩', '-120', '+120'],
    ['J3', '大臂横滚', '-170', '+170'],
    ['J4', '肘', '-120', '+120'],
    ['J5', '小臂横滚', '-170', '+170'],
    ['J6', '腕', '-120', '+120'],
    ['J7', '末端横滚', '-175', '+175'],
]
for j, h in enumerate(headers3):
    cell = table3.rows[0].cells[j]
    cell.text = h
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.bold = True
            r.font.size = Pt(9)
for i, row_data in enumerate(data3):
    for j, val in enumerate(row_data):
        cell = table3.rows[i+1].cells[j]
        cell.text = val
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.size = Pt(9)

add_para('', space_after=8)

add_heading('4.3  正运动学测试', level=2)

add_para('测试2验证零位构型（所有关节角为0°）下的正运动学精度：在零位构型下，'
         '末端执行器的X和Y坐标应为0（位于世界坐标系原点正上方），Z坐标应约为0.87米。'
         '实际测试结果为末端位置（0.000, 0.000, 0.870）米，与理论设计值吻合。', first_line_indent=0.74, space_after=6)

add_para('测试3验证随机构型下的FK一致性：设随机种子rng(123)以保证可复现性，'
         '对10组随机生成的构型分别计算末端位姿的齐次变换矩阵，验证其旋转矩阵部分的行列式为1（即旋转矩阵正交、有效）。'
         '10组随机测试全部通过，确认正运动学在各构型下均产生正确的刚体变换。', first_line_indent=0.74, space_after=8)

add_image('正运动学演示.gif', width_inches=5.0, caption='图4-1  正运动学演示——关节角度变化与末端轨迹')

add_heading('4.4  逆运动学测试', level=2)

add_para('测试5全面验证逆运动学求解的正确性和鲁棒性，分为两个子测试：', first_line_indent=0.74, space_after=4)

add_para('子测试1：零位自洽性验证。以零位构型为起点，将零位构型对应的末端位置作为IK目标，'
         '验证求解器能否恢复出零位关节角（即各关节角接近0）。测试结果：求解成功（Status="success"），'
         '解q_sol的范数||q_sol|| < 10⁻³，确认了求解器的自洽性。', first_line_indent=0.74, space_after=4)

add_para('子测试2：随机目标验证。设随机种子rng(456)，生成5组随机目标位置。'
         '每个目标位置的生成方式为：取随机构型qRand的末端位置，'
         '在X方向微调+0.02米，Z方向限制在[0.3, 0.8]米的工作空间范围内，然后以零位为初始猜测进行IK求解。'
         '验证标准为：①求解状态为"success"；②求解得到的末端位置与目标位置之间的欧氏距离误差 < 0.01米。'
         '5组随机目标全部通过验证，逆运动学求解成功率达到100%。', first_line_indent=0.74, space_after=8)

add_heading('4.5  轨迹规划测试', level=2)

add_para('测试8验证轨迹生成的一致性：以零位构型作为起始和结束路径点，中间加入一个偏移点（+0.2m X方向，+0.1m Z方向），'
         '构成三路径点序列[home, offset, home]。使用梯形速度剖面法在2秒内生成50个采样点的轨迹。'
         '验证标准为：①轨迹矩阵尺寸为6×50（6个关节×50个时间步）；'
         '②轨迹第一帧的关节角等于起始构型（||qTraj(:,1) - home|| < 10⁻⁶）；'
         '③轨迹最后一帧的关节角等于结束构型（||qTraj(:,end) - home|| < 10⁻⁶），'
         '即机械臂完成从home出发、经过中间点、再返回home的完整往返运动。测试结果符合预期。', first_line_indent=0.74, space_after=6)

add_para('此外，测试7验证了零位构型的自碰撞安全性：使用checkCollision函数（跳过父子相邻连杆的碰撞检测）'
         '检查零位构型是否发生自碰撞，结果确认零位构型下各连杆之间无碰撞，机械臂结构设计合理。', first_line_indent=0.74, space_after=6)

# Test results summary table
add_para('表4-2  自动化测试结果汇总', bold=True, size=10,
         alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
table4 = doc.add_table(rows=9, cols=3, style='Table Grid')
table4.alignment = WD_ALIGN_PARAGRAPH.CENTER
headers4 = ['测试编号', '测试项目', '结果']
data4 = [
    ['Test 1', '模型结构完整性', 'PASS（9个刚体，名称正确）'],
    ['Test 2', '零位FK精度', 'PASS（EE位置 0,0,0.87）'],
    ['Test 3', '随机FK一致性（10组）', 'PASS（det(R)=1）'],
    ['Test 4', '雅可比矩阵秩', 'PASS（6×7，rank=6）'],
    ['Test 5', '逆运动学求解（6组）', 'PASS（成功率100%）'],
    ['Test 6', '关节限位合理性', 'PASS（6个关节均有效）'],
    ['Test 7', '零位自碰撞检测', 'PASS（无碰撞）'],
    ['Test 8', '轨迹生成一致性', 'PASS（起止匹配）'],
]
for j, h in enumerate(headers4):
    cell = table4.rows[0].cells[j]
    cell.text = h
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.bold = True
            r.font.size = Pt(9)
for i, row_data in enumerate(data4):
    for j, val in enumerate(row_data):
        cell = table4.rows[i+1].cells[j]
        cell.text = val
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.size = Pt(9)

add_para('', space_after=6)
add_para('全部8项测试均通过，系统功能和精度满足设计要求。', bold=True, first_line_indent=0.74, space_after=8)

add_page_break()

# ╔══════════════════════════════════════════════════════════════╗
# ║  CHAPTER 5 — 总结与展望                                      ║
# ╚══════════════════════════════════════════════════════════════╝

add_heading('5  总结与展望', level=1)

add_heading('5.1  工作总结', level=2)

add_para('本项目基于MATLAB Robotics System Toolbox，设计并实现了一套功能完整的七自由度机械臂运动学仿真与轨迹规划系统，'
         '主要完成了以下工作：', first_line_indent=0.74, space_after=6)

summary_items = [
    ('建立了精确的七自由度机械臂刚体树模型：', '包含底座、七段旋转连杆和末端执行器共9个刚体，'
     '每段连杆配置了合理的几何尺寸参数和关节运动范围限制，采用Z-Y-Z-Y-Z-Y-Z关节构型，'
     '通过R_x(±90°)固定变换实现坐标系交替转换，零位构型下末端高度0.87米。'),
    ('实现了完整的运动学分析体系：', '正运动学支持任意构型下末端位姿的精确计算；'
     '逆运动学基于GIK数值求解器，支持仅位置约束和位姿全约束两种模式，在测试中达到100%求解成功率；'
     '雅可比矩阵分析验证了机械臂在非奇异构型下的满秩特性。'),
    ('集成了四种轨迹规划方法：', '梯形速度剖麵法、三次多项式、五次多项式和B样条曲线，'
     '支持用户自定义路径点序列和运动参数，自动完成IK求解、轨迹生成、三维动画播放和关节曲线绘制。'),
    ('开发了功能完善的交互式GUI：', '基于MATLAB uifigure框架，采用双标签页设计（探索+规划），'
     '提供FK/IK滑块交互探索和完整的轨迹规划工作流。界面采用防抖定时器、动画代数标记等技术确保流畅的用户体验。'),
    ('编写了系统化的测试脚本：', '8项自动化测试覆盖模型结构、FK、IK、雅可比、碰撞和轨迹各核心模块，'
     '全部测试通过，验证了系统的正确性和可靠性。'),
]
for title, text in summary_items:
    add_para(title + text, first_line_indent=0.74, space_after=5)

add_para('', space_after=4)
add_para('通过本项目，深入理解了串联机械臂的运动学原理、数值逆解方法、关节空间轨迹规划算法以及MATLAB机器人工具箱的使用方法，'
         '为后续学习机器人动力学、运动控制和实际机器人系统开发奠定了坚实的基础。', first_line_indent=0.74, space_after=8)

add_heading('5.2  不足与展望', level=2)

add_para('本项目虽然实现了预期的全部功能，但仍存在以下不足和可改进之处：', first_line_indent=0.74, space_after=6)

limitations = [
    ('动力学建模缺失：', '当前系统仅覆盖运动学层面的仿真，未包含动力学建模（质量、惯量、重力补偿、力矩计算等）。'
     '未来可基于RigidBodyTree的动力学属性扩展，实现基于力矩的前馈控制和阻抗控制仿真。'),
    ('避障功能未实现：', '虽然七自由度机械臂的核心优势在于冗余自由度带来的避障能力，但当前系统尚未集成碰撞检测与避障路径规划功能。'
     '未来可结合MATLAB的碰撞检测API（checkCollision）和路径规划算法（如RRT、PRM），实现自动避障轨迹规划。'),
    ('轨迹规划仅限关节空间：', '当前四种轨迹方法均在关节空间中生成轨迹，虽然保证了关节运动的平滑性，'
     '但末端执行器的笛卡尔路径形状不受直接控制。未来可增加笛卡尔空间轨迹规划（如直线、圆弧插补），'
     '以满足焊接、涂胶等对末端路径形状有严格要求的应用场景。'),
    ('GUI交互可进一步优化：', '当前GUI基于MATLAB uifigure开发，功能和美观性受限于MATLAB图形框架。'
     '未来可考虑基于MATLAB App Designer重新设计界面，添加更多可视化元素（如速度仪表盘、工作空间云图），'
     '或通过MATLAB与ROS的接口将仿真系统连接到真实的机器人控制器。'),
    ('缺少与真机联调：', '当前系统为纯仿真环境，尚未与实际机械臂硬件进行联调验证。'
     '未来可以输出关节轨迹数据，通过串口/以太网通信发送给真实的七轴机械臂控制器，实现仿真到实机的闭环。'),
]
for title, text in limitations:
    add_para(title + text, first_line_indent=0.74, space_after=5)

add_para('', space_after=4)
add_para('总之，本项目为七自由度机械臂的运动学仿真提供了一个较为完整的基础框架，'
         '在此基础上可以方便地扩展动力学、避障、笛卡尔规划和真机联调等功能，'
         '具有较高的教学参考价值和工程应用前景。', first_line_indent=0.74, space_after=8)

add_page_break()

# ╔══════════════════════════════════════════════════════════════╗
# ║  REFERENCES (参考文献)                                       ║
# ╚══════════════════════════════════════════════════════════════╝

add_heading('参考文献', level=1)

references = [
    '[1] MathWorks. Robotics System Toolbox Documentation: Rigid Body Tree Robot Models[EB/OL]. '
    'https://www.mathworks.com/help/robotics/rigid-body-tree-robot-models.html, 2024.',
    '[2] MathWorks. Robotics System Toolbox Documentation: Inverse Kinematics[EB/OL]. '
    'https://www.mathworks.com/help/robotics/inverse-kinematics.html, 2024.',
    '[3] MathWorks. Robotics System Toolbox Documentation: Trajectory Generation for Robot Manipulators[EB/OL]. '
    'https://www.mathworks.com/help/robotics/trajectory-generation.html, 2024.',
    '[4] Corke P. Robotics, Vision and Control: Fundamental Algorithms in MATLAB[M]. '
    '2nd ed. Berlin: Springer, 2017.',
    '[5] Siciliano B, Sciavicco L, Villani L, et al. Robotics: Modelling, Planning and Control[M]. '
    'London: Springer, 2009.',
    '[6] Craig J J. Introduction to Robotics: Mechanics and Control[M]. 4th ed. Harlow: Pearson, 2017.',
    '[7] Lynch K M, Park F C. Modern Robotics: Mechanics, Planning, and Control[M]. '
    'Cambridge: Cambridge University Press, 2017.',
    '[8] Spong M W, Hutchinson S, Vidyasagar M. Robot Modeling and Control[M]. '
    '2nd ed. Hoboken: John Wiley & Sons, 2020.',
    '[9] 蔡自兴. 机器人学[M]. 3版. 北京: 清华大学出版社, 2015.',
    '[10] 熊有伦, 丁汉, 刘恩沧. 机器人学: 建模、控制与视觉[M]. 2版. 武汉: 华中科技大学出版社, 2018.',
    '[11] 霍伟. 机器人动力学与控制[M]. 北京: 高等教育出版社, 2005.',
    '[12] Featherstone R. Rigid Body Dynamics Algorithms[M]. New York: Springer, 2008.',
]
for ref in references:
    add_para(ref, space_after=3, size=10.5)

add_page_break()

# ╔══════════════════════════════════════════════════════════════╗
# ║  ACKNOWLEDGEMENTS (致谢)                                    ║
# ╚══════════════════════════════════════════════════════════════╝

add_heading('致  谢', level=1)

add_para('在本项目的设计与实现过程中，得到了指导老师的悉心指导和诸多帮助。老师在教学过程中深入浅出地讲解了'
         '机器人运动学的基本原理和MATLAB Robotics System Toolbox的使用方法，为本项目的顺利开展提供了坚实的理论基础。'
         '同时，感谢团队成员在项目讨论中提出的宝贵意见和建议，使系统设计更加完善。', first_line_indent=0.74, space_after=6)

add_para('此外，MathWorks公司提供的Robotics System Toolbox文档和示例代码为本项目提供了重要的技术参考，'
         '开源社区的丰富资源也为项目的开发带来了诸多启发。在此一并表示诚挚的谢意。', first_line_indent=0.74, space_after=12)

add_para('', space_after=24)
# Signature line
p_sig = doc.add_paragraph()
p_sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run_sig = p_sig.add_run('谢毅捷\n2026年7月2日')
set_run_font(run_sig, '宋体', 12)

# ═══════════════════════════════════════════════════════════════
#  SAVE
# ═══════════════════════════════════════════════════════════════
doc.save(OUTPUT)
print(f'Document saved to: {OUTPUT}')
print('Done!')
